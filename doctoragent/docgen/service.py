"""Document export (chat → Markdown / PDF / DOCX).

Turns a chat session's messages into a downloadable document:
* Markdown (``.md``)
* PDF (``.pdf``) via reportlab
* Word (``.docx``) via python-docx

Also provides :func:`render_markdown` for converting arbitrary markdown text
to PDF/DOCX so the assistant can turn an answer into a shareable file.
"""

from __future__ import annotations

import io
from collections.abc import Iterable
from pathlib import Path
from typing import Any

# reportlab / python-docx are imported lazily inside markdown_to_pdf / _docx
# and degrade gracefully when absent; the `server` extra ships both.

_SUPPORTED = ("md", "pdf", "docx")


class DocExportError(RuntimeError):
    pass


def _md_escape(text: str) -> str:
    return (text or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def messages_to_markdown(messages: Iterable[dict[str, Any]]) -> str:
    """Render chat messages (role/content) as a Markdown transcript."""
    lines: list[str] = ["# DoctorAgent 对话记录", ""]
    for m in messages:
        role = (m.get("role") or "user").upper()
        content = m.get("content", "")
        if not isinstance(content, str):
            content = str(content or "")
        lines.append(f"## {role}")
        lines.append("")
        lines.append(content)
        lines.append("")
    return "\n".join(lines)


def markdown_to_pdf(md: str, out: Path) -> Path:
    """Render a Markdown document to PDF via reportlab (simple heading/para)."""
    # Re-check importability at call time so installing reportlab later (or a
    # stale import-time flag) does not block export. reportlab is a runtime
    # dependency of the `server` extra, so it is normally always available.
    try:
        from reportlab.lib.pagesizes import A4  # noqa: PLC0415
        from reportlab.lib.styles import ParagraphStyle  # noqa: PLC0415
        from reportlab.lib.units import mm  # noqa: PLC0415
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer  # noqa: PLC0415
    except ImportError:
        raise DocExportError("PDF export requires reportlab (pip install reportlab)") from None
    styles = {
        "h1": ParagraphStyle(
            "h1", fontName="Helvetica-Bold", fontSize=20, leading=26, spaceAfter=10
        ),
        "h2": ParagraphStyle(
            "h2", fontName="Helvetica-Bold", fontSize=15, leading=20, spaceAfter=8, spaceBefore=10
        ),
        "body": ParagraphStyle("body", fontName="Helvetica", fontSize=10.5, leading=15),
    }
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=18 * mm,
        rightMargin=18 * mm,
        topMargin=18 * mm,
        bottomMargin=18 * mm,
    )
    story: list[Any] = []
    for raw in md.splitlines():
        line = raw.strip()
        if not line:
            continue
        if line.startswith("### "):
            story.append(Paragraph(_md_escape(line[4:]), styles["h2"]))
        elif line.startswith("## "):
            story.append(Paragraph(_md_escape(line[3:]), styles["h1"]))
        elif line.startswith("# "):
            story.append(Paragraph(_md_escape(line[2:]), styles["h1"]))
        else:
            story.append(Paragraph(_md_escape(line), styles["body"]))
        story.append(Spacer(1, 4))
    doc.build(story)
    out.write_bytes(buf.getvalue())
    return out


def markdown_to_docx(md: str, out: Path) -> Path:
    """Render a Markdown document to .docx via python-docx."""
    try:
        from docx import Document  # noqa: PLC0415
    except ImportError:
        raise DocExportError("DOCX export requires python-docx (pip install python-docx)") from None
    doc = Document()
    for raw in md.splitlines():
        line = raw.strip()
        if not line:
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("- ") or line.startswith("* "):
            doc.add_paragraph(line[2:], style="List Bullet")
        else:
            doc.add_paragraph(line)
    doc.save(str(out))
    return out


def export_messages(messages: Iterable[dict[str, Any]], fmt: str, out: Path) -> Path:
    """Export chat messages to the requested format (md | pdf | docx)."""
    fmt = (fmt or "md").lower().lstrip(".")
    if fmt not in _SUPPORTED:
        raise DocExportError(f"unsupported format {fmt!r}; choose from {_SUPPORTED}")
    md = messages_to_markdown(messages)
    out.parent.mkdir(parents=True, exist_ok=True)
    if fmt == "md":
        out.write_text(md, encoding="utf-8")
        return out
    if fmt == "pdf":
        return markdown_to_pdf(md, out)
    return markdown_to_docx(md, out)
